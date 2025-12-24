from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from .models import ResearchSession, UploadedDocument, ResearchSummary, ResearchReasoning, ResearchCost
from django.shortcuts import get_object_or_404
from django.contrib.auth.models import User
from django.http import JsonResponse

# Home page
def home(request):
    return JsonResponse({"message": "Welcome to the Research API"})


# Start a new research
class StartResearch(APIView):
    def post(self, request):
        user = User.objects.first()  # Replace with request.user when auth is added
        query = request.data.get("query")
        parent_id = request.data.get("parent_research_id")
        parent = get_object_or_404(ResearchSession, id=parent_id) if parent_id else None

        research = ResearchSession.objects.create(
            user=user,
            query=query,
            parent=parent,
            status="STARTED"
        )
        ResearchSummary.objects.create(research=research)
        ResearchReasoning.objects.create(research=research)
        ResearchCost.objects.create(research=research)

        return Response({"message": "Research started", "id": research.id}, status=status.HTTP_201_CREATED)


# Continue research
class ContinueResearch(APIView):
    def post(self, request, research_id):
        parent_research = get_object_or_404(ResearchSession, id=research_id)
        query = request.data.get("query")

        new_research = ResearchSession.objects.create(
            user=parent_research.user,
            query=query,
            parent=parent_research,
            status="STARTED"
        )
        ResearchSummary.objects.create(research=new_research)
        ResearchReasoning.objects.create(research=new_research)
        ResearchCost.objects.create(research=new_research)

        return Response({"message": "Research continued", "id": new_research.id}, status=status.HTTP_201_CREATED)


# Upload document with text extraction
class UploadDocument(APIView):
    def post(self, request, research_id):
        research = get_object_or_404(ResearchSession, id=research_id)
        file = request.FILES.get("file")
        if not file:
            return Response({"error": "No file uploaded"}, status=status.HTTP_400_BAD_REQUEST)
        
        # Save uploaded file
        doc = UploadedDocument.objects.create(
            research=research,
            file=file
        )

       
        import PyPDF2
        from docx import Document

        extracted_text = ""
        file_path = doc.file.path

        try:
            if file.name.endswith(".pdf"):
                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        extracted_text += page.extract_text() or ""
            elif file.name.endswith(".txt"):
                with open(file_path, "r", encoding="utf-8") as f:
                    extracted_text = f.read()
            elif file.name.endswith(".docx"):
                docx_file = Document(file_path)
                extracted_text = "\n".join([p.text for p in docx_file.paragraphs])
            else:
                extracted_text = "(unsupported file type, cannot extract text)"
        except Exception as e:
            extracted_text = f"(error extracting text: {str(e)})"
        # -------------------------------

        doc.extracted_text = extracted_text
        doc.save()

        return Response({
            "message": "File uploaded and extracted",
            "file_id": doc.id,
            "extracted_text_preview": extracted_text[:300]  # first 300 chars
        }, status=status.HTTP_201_CREATED)


# Research history
class ResearchHistory(APIView):
    def get(self, request):
        user = User.objects.first()
        researches = ResearchSession.objects.filter(user=user)
        data = [{"id": r.id, "query": r.query, "status": r.status} for r in researches]
        return Response(data)


# Research details
class ResearchDetails(APIView):
    def get(self, request, research_id):
        research = get_object_or_404(ResearchSession, id=research_id)
        summary = ResearchSummary.objects.get(research=research).summary
        reasoning = ResearchReasoning.objects.get(research=research).reasoning
        cost = ResearchCost.objects.get(research=research)
        documents = [{"id": d.id, "filename": d.file.name, "extracted_text_preview": d.extracted_text[:300]} 
                     for d in research.documents.all()]

        return Response({
            "id": research.id,
            "query": research.query,
            "status": research.status,
            "summary": summary,
            "reasoning": reasoning,
            "cost": {
                "input_tokens": cost.input_tokens,
                "output_tokens": cost.output_tokens,
                "total_cost": cost.total_cost
            },
            "documents": documents
        })
